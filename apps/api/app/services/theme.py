"""Report theming: theme color + page background templates.

Applied at render time on the generated DOCX:
  - theme_color: recolors every element that uses the template's accent
    (sentinel) color to the user's chosen hex, and sets the document page
    background color for solid backgrounds.
  - background: applies a page background from the bundled catalog (solid
    color via w:background, or an image via VML background in settings.xml),
    or a custom uploaded image.

Everything is validated; invalid input falls back to defaults so rendering
never fails.
"""

import io
import re

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import RGBColor

# The accent color the templates are built with (a.k.a. the recolor sentinel).
SENTINEL = RGBColor(0x0B, 0x6E, 0x6B)
SENTINEL_HEX = "0B6E6B"

HEX_RE = re.compile(r"^#?([0-9A-Fa-f]{6})$")

# Background catalog: id -> {name, kind, color?, stops?}
# kind: "none" | "solid" | "gradient" | "pattern"
BACKGROUNDS = {
    "none": {"name": "None", "kind": "none"},
    "solid-cream": {"name": "Cream", "kind": "solid", "color": "FAF6EC"},
    "solid-mist": {"name": "Mist", "kind": "solid", "color": "EFF5F5"},
    "solid-sand": {"name": "Sand", "kind": "solid", "color": "F6F1E7"},
    "solid-blush": {"name": "Blush", "kind": "solid", "color": "FDF0EE"},
    "solid-sage": {"name": "Sage", "kind": "solid", "color": "EFF3EA"},
    "solid-ink": {"name": "Ink", "kind": "solid", "color": "1B2A41"},
    "grad-teal": {"name": "Teal Gradient", "kind": "gradient", "stops": ["0B6E6B", "FFFFFF"]},
    "grad-ocean": {"name": "Ocean Gradient", "kind": "gradient", "stops": ["1B4965", "F2F6FB"]},
    "grad-forest": {"name": "Forest Gradient", "kind": "gradient", "stops": ["2F4F2F", "F6F1E7"]},
    "grad-sunset": {"name": "Sunset Gradient", "kind": "gradient", "stops": ["C97B63", "FDEBDD"]},
    "grad-lilac": {"name": "Lilac Gradient", "kind": "gradient", "stops": ["7C6FA0", "F3F0FA"]},
    "pat-dots": {"name": "Dots", "kind": "pattern", "color": "D8E7E6"},
    "pat-stripes": {"name": "Stripes", "kind": "pattern", "color": "DCE9E8"},
    "pat-grid": {"name": "Grid", "kind": "pattern", "color": "D8E4E3"},
}


def parse_hex(value) -> str | None:
    """Return a 6-char uppercase hex (no '#') or None when invalid."""
    if not value:
        return None
    m = HEX_RE.match(str(value).strip())
    if not m:
        return None
    return m.group(1).upper()


def background_kind(bg_id) -> str | None:
    if not bg_id:
        return None
    return (BACKGROUNDS.get(bg_id) or {}).get("kind")


def render_background_image(bg_id: str, width: int, height: int) -> bytes:
    """Generate a background image (gradient or pattern) as PNG bytes."""
    from PIL import Image, ImageDraw

    meta = BACKGROUNDS[bg_id]
    img = Image.new("RGB", (width, height), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    if meta["kind"] == "gradient":
        stops = meta["stops"]
        top = tuple(int(stops[0][i : i + 2], 16) for i in (0, 2, 4))
        bottom = tuple(int(stops[1][i : i + 2], 16) for i in (0, 2, 4))
        for y in range(height):
            t = y / max(height - 1, 1)
            color = tuple(int(top[c] + (bottom[c] - top[c]) * t) for c in range(3))
            draw.line([(0, y), (width, y)], fill=color)
    elif meta["kind"] == "pattern":
        base = tuple(int(meta["color"][i : i + 2], 16) for i in (0, 2, 4))
        img.paste(base, (0, 0, width, height))
        accent = (255, 255, 255)
        if meta["color"] == "D8E7E6":  # dots
            r = 10
            for x in range(0, width, 46):
                for y in range(0, height, 46):
                    draw.ellipse([x, y, x + 2 * r, y + 2 * r], fill=accent)
        elif meta["color"] == "DCE9E8":  # diagonal stripes
            for offset in range(-height, width, 48):
                draw.line([(offset, height), (offset + height, 0)], fill=accent, width=14)
        else:  # grid
            for x in range(0, width, 64):
                draw.line([(x, 0), (x, height)], fill=accent, width=2)
            for y in range(0, height, 64):
                draw.line([(0, y), (width, y)], fill=accent, width=2)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _recolor(doc: Document, color: RGBColor) -> None:
    """Replace the template accent color with the chosen theme color."""
    def _runs(paragraph):
        for run in paragraph.runs:
            if run.font.color and run.font.color.type is not None:
                try:
                    if run.font.color.rgb == SENTINEL:
                        run.font.color.rgb = color
                except (TypeError, ValueError):
                    continue

    for paragraph in doc.paragraphs:
        _runs(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _runs(paragraph)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                _runs(paragraph)


def _set_solid_background(doc: Document, hex_color: str) -> None:
    settings = doc.settings.element
    for bg in settings.findall(qn("w:background")):
        settings.remove(bg)
    bg = parse_xml(
        f'<w:background {nsdecls("w")} w:color="{hex_color}"/>'
    )
    settings.append(bg)


def _set_image_background(doc: Document, image_bytes: bytes) -> None:
    settings = doc.settings.element
    for bg in settings.findall(qn("w:background")):
        settings.remove(bg)

    r_id, _ = doc.part.get_or_add_image(io.BytesIO(image_bytes))
    vml = (
        f'<w:background {nsdecls("w")} '
        f'xmlns:v="urn:schemas-microsoft-com:vml" '
        f'xmlns:o="urn:schemas-microsoft-com:office:office" w:color="FFFFFF">'
        f'<w:drawing>'
        f'<v:background id="_x0000_s1025" o:bwmode="auto">'
        f'<v:fill o:relid="{r_id}" type="frame"/>'
        f"</v:background>"
        f"</w:drawing>"
        f"</w:background>"
    )
    settings.append(parse_xml(vml))


def background_png_bytes(bg_id: str, width: int, height: int) -> bytes | None:
    """Return a PNG for any background id (solids/gradients/patterns), or None
    for 'none'/unknown ids."""
    meta = BACKGROUNDS.get(bg_id)
    if meta is None or meta["kind"] == "none":
        return None
    if meta["kind"] == "solid":
        from PIL import Image

        color = tuple(int(meta["color"][i : i + 2], 16) for i in (0, 2, 4))
        img = Image.new("RGB", (width, height), color)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    return render_background_image(bg_id, width, height)


def apply_background_to_pdf(pdf_bytes: bytes, image_bytes: bytes) -> bytes:
    """Overlay a background image behind the content of every PDF page.

    LibreOffice's DOCX import does not reliably render page backgrounds, so the
    background is applied here as a guaranteed post-processing step.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page in doc:
        page.insert_image(page.rect, stream=image_bytes, overlay=False)
    out = io.BytesIO()
    doc.save(out, garbage=3, deflate=True)
    doc.close()
    return out.getvalue()


def apply_theme(
    docx_bytes: bytes,
    theme_color: str | None = None,
    background_id: str | None = None,
    custom_background: bytes | None = None,
) -> bytes:
    """Recolor the accent to theme_color and apply the page background.

    Always succeeds: invalid theme values are ignored (document unchanged for
    that aspect) rather than raising.
    """
    doc = Document(io.BytesIO(docx_bytes))

    color_hex = parse_hex(theme_color)
    if color_hex and color_hex != SENTINEL_HEX:
        _recolor(doc, RGBColor.from_string(color_hex))

    if custom_background:
        _set_image_background(doc, custom_background)
    else:
        kind = background_kind(background_id)
        if kind == "solid":
            _set_solid_background(doc, BACKGROUNDS[background_id]["color"])
        elif kind in ("gradient", "pattern"):
            _set_image_background(doc, render_background_image(background_id, 1240, 1754))

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()