"""Rich-text section overlay.

When a user edits a section in the Word-like editor (TipTap HTML), the rebuild
folds the edited content back into the DOCX with formatting (bold, italic,
underline, strike, color, highlight, alignment, headings, lists, quotes).

The template placeholder paragraph (rendered to the section's plain text) is
replaced with rich paragraphs built from the HTML. Any parsing problem falls
back to the plain text so generation never breaks.
"""

import io
import re
from html.parser import HTMLParser

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

_BLOCK_TAGS = {"p", "div", "h1", "h2", "h3", "h4", "h5", "li", "blockquote", "br"}
_INLINE_BOLD = {"b", "strong"}
_INLINE_ITALIC = {"i", "em"}
_INLINE_UNDERLINE = {"u", "ins"}
_INLINE_STRIKE = {"s", "del", "strike"}


class _Run:
    __slots__ = ("text", "bold", "italic", "underline", "strike", "color", "highlight")

    def __init__(self, text="", bold=False, italic=False, underline=False, strike=False, color=None, highlight=None):
        self.text = text
        self.bold = bold
        self.italic = italic
        self.underline = underline
        self.strike = strike
        self.color = color
        self.highlight = highlight


class _Block:
    __slots__ = ("kind", "align", "runs", "list_type")

    def __init__(self, kind="p", align=None, runs=None, list_type=None):
        self.kind = kind
        self.align = align
        self.runs = runs or []
        self.list_type = list_type


class _Converter(HTMLParser):
    def __init__(self):
        super().__init__()
        self.blocks: list[_Block] = []
        self._stack: list[dict] = [{}]  # formatting context stack
        self._cur: list[_Run] | None = None
        self._list_stack: list[str] = []
        self._pending_br = False

    def _fmt(self):
        f = {}
        for ctx in self._stack:
            f.update(ctx)
        return f

    def _flush_block(self, list_type=None):
        top = self._stack[0]
        if list_type:
            kind = "li"
        else:
            kind = top.get("heading") or ("blockquote" if top.get("quote") else "p")
        if self._cur and any(r.text for r in self._cur):
            self.blocks.append(
                _Block(kind=kind, align=top.get("align"), runs=self._cur, list_type=list_type)
            )
        self._cur = None

    def handle_starttag(self, tag, attrs):
        attrs = dict(attrs)
        if tag in _BLOCK_TAGS:
            self._flush_block()
            self._cur = []
            align = None
            style = attrs.get("style", "")
            m = re.search(r"text-align:\s*(left|center|right|justify)", style)
            if m:
                align = m.group(1)
            if tag in ("h1", "h2", "h3", "h4", "h5"):
                self._stack[0]["heading"] = tag
            elif tag == "blockquote":
                self._stack[0]["quote"] = True
            if align:
                self._stack[0]["align"] = align
            return
        if tag in ("ul", "ol"):
            self._list_stack.append("bullet" if tag == "ul" else "number")
            return
        if tag == "li":
            self._flush_block(list_type=self._list_stack[-1] if self._list_stack else "bullet")
            self._cur = []
            return
        if tag == "br":
            if self._cur:
                self._cur.append(_Run(text="\n"))
            return
        ctx = {}
        if tag in _INLINE_BOLD:
            ctx["bold"] = True
        elif tag in _INLINE_ITALIC:
            ctx["italic"] = True
        elif tag in _INLINE_UNDERLINE:
            ctx["underline"] = True
        elif tag in _INLINE_STRIKE:
            ctx["strike"] = True
        elif tag == "span":
            style = attrs.get("style", "")
            m = re.search(r"color:\s*(#[0-9a-fA-F]{6}|rgb\([^)]*\))", style)
            if m:
                val = m.group(1)
                if val.startswith("rgb"):
                    parts = re.findall(r"\d+", val)
                    if len(parts) == 3:
                        val = "#{:02X}{:02X}{:02X}".format(*[int(p) for p in parts])
                ctx["color"] = val
            m = re.search(r"background-color:\s*(#[0-9a-fA-F]{6}|rgb\([^)]*\))", style)
            if m:
                val = m.group(1)
                if val.startswith("rgb"):
                    parts = re.findall(r"\d+", val)
                    if len(parts) == 3:
                        val = "#{:02X}{:02X}{:02X}".format(*[int(p) for p in parts])
                ctx["highlight"] = val
        if ctx:
            self._stack.append(ctx)
        else:
            self._stack.append({})

    def handle_endtag(self, tag):
        if tag in _BLOCK_TAGS:
            list_type = self._list_stack[-1] if (self._list_stack and tag == "li") else None
            self._flush_block(list_type=list_type)
            if tag in ("h1", "h2", "h3", "h4", "h5"):
                self._stack[0].pop("heading", None)
            elif tag == "blockquote":
                self._stack[0].pop("quote", None)
            self._stack[0].pop("align", None)
            return
        if tag in ("ul", "ol"):
            if self._list_stack:
                self._list_stack.pop()
            return
        if len(self._stack) > 1:
            self._stack.pop()

    def handle_data(self, data):
        if self._cur is None:
            self._cur = []
        f = self._fmt()
        self._cur.append(
            _Run(
                text=data,
                bold=f.get("bold", False),
                italic=f.get("italic", False),
                underline=f.get("underline", False),
                strike=f.get("strike", False),
                color=f.get("color"),
                highlight=f.get("highlight"),
            )
        )


def _parse_html(html: str) -> list[_Block]:
    parser = _Converter()
    parser.feed(html or "")
    parser.close()
    parser._flush_block()
    return parser.blocks


_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _hex_to_rgb(hex_color: str | None) -> RGBColor | None:
    if not hex_color:
        return None
    m = re.match(r"^#?([0-9A-Fa-f]{6})$", hex_color.strip())
    if not m:
        return None
    try:
        return RGBColor.from_string(m.group(1))
    except (ValueError, TypeError):
        return None


def _style_for(block: _Block):
    if block.kind.startswith("h"):
        level = int(block.kind[1]) if block.kind[1:].isdigit() else 1
        return f"Heading {min(level, 4)}", 12 + (4 - min(level, 4)) * 2, True
    if block.list_type == "bullet":
        return "List Bullet", 11, False
    if block.list_type == "number":
        return "List Number", 11, False
    if block.kind == "blockquote":
        return None, 11, False
    return None, 11, False


def _add_block(doc: Document, block: _Block) -> None:
    style_name, size, bold_all = _style_for(block)
    try:
        p = doc.add_paragraph(style=style_name) if style_name else doc.add_paragraph()
    except KeyError:
        p = doc.add_paragraph()
    if block.align and block.align in _ALIGN:
        p.alignment = _ALIGN[block.align]
    if block.kind == "blockquote":
        p.paragraph_format.left_indent = Pt(24)
    for run in block.runs:
        text = run.text
        if not text:
            continue
        r = p.add_run(text)
        r.bold = run.bold or bold_all
        r.italic = run.italic or block.kind == "blockquote"
        r.underline = run.underline
        r.strike = run.strike
        if size:
            r.font.size = Pt(size)
        color = _hex_to_rgb(run.color)
        if color:
            r.font.color.rgb = color


def replace_placeholder_with_rich(doc: Document, placeholder_text: str, html: str) -> bool:
    """Replace the paragraph matching placeholder_text with rich paragraphs."""
    target = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == (placeholder_text or "").strip():
            target = paragraph
            break
    if target is None:
        return False

    blocks = _parse_html(html)
    if not blocks:
        return False

    # Build new paragraphs after the target, then remove the target.
    anchor = target._p
    new_paragraphs = []
    for block in blocks:
        _add_block(doc, block)
        new_paragraphs.append(doc.paragraphs[-1])
    for p in new_paragraphs:
        anchor.addnext(p._p)
        anchor = p._p
    target._p.getparent().remove(target._p)
    return True


def apply_rich_sections(docx_bytes: bytes, overrides: list[tuple[str, str]]) -> bytes:
    """Apply section HTML overrides to a rendered DOCX.

    overrides = [(placeholder_plain_text, content_html)]. Any failure is
    silently skipped (the plain text placeholder remains) so generation is safe.
    """
    if not overrides:
        return docx_bytes
    doc = Document(io.BytesIO(docx_bytes))
    for plain, html in overrides:
        try:
            replace_placeholder_with_rich(doc, plain, html)
        except Exception:
            continue
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()