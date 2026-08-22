"""Generation pipeline services.

Core flow:
  1. Load template .docx from storage.
  2. docxtpl.render() replaces {{ var }} placeholders (jinja2 syntax).
  3. python-docx swaps [img:NAME] markers for uploaded / placeholder images.
  4. Save editable .docx to storage.
  5. LibreOffice headless converts DOCX -> PDF.
  6. mammoth converts DOCX -> HTML for in-browser preview.

DocxTemplate never executes macros and runs with no network, so macro/XXE risks
are avoided (verified against official docxtpl docs).
"""

import io
import re
import subprocess
import tempfile
import os
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docxtpl import DocxTemplate

from ..storage import ObjectData

IMAGE_MARKER_RE = re.compile(r"\[img:([A-Za-z0-9_.\-]+)\]")

# Sensible default widths (inches) per image marker name.
DEFAULT_IMAGE_WIDTHS = {
    "logo": 1.5,
    "hero": 6.0,
    "chart_funding": 5.5,
    "chart": 5.5,
}


class TemplateValidationError(ValueError):
    """Raised when an uploaded template is not a valid .docx."""


def validate_docx(data: bytes) -> None:
    """Validate that bytes are a real .docx (zip containing [Content_Types].xml)."""
    try:
        import zipfile

        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            names = zf.namelist()
    except zipfile.BadZipFile as exc:
        raise TemplateValidationError("File is not a valid DOCX (not a zip)") from exc
    if "[Content_Types].xml" not in names:
        raise TemplateValidationError("File is not a valid DOCX (missing content types)")


def html_to_plain_text(html: str | None) -> str:
    """Strip HTML tags to plain text (stdlib only, safe for untrusted HTML)."""
    if not html:
        return ""
    import html as html_lib

    from html.parser import HTMLParser

    class _TextExtractor(HTMLParser):
        def __init__(self):
            super().__init__()
            self.parts: list[str] = []
            self.skip = 0

        def handle_starttag(self, tag, attrs):
            if tag in ("script", "style"):
                self.skip += 1
            if tag in ("p", "div", "br", "li", "h1", "h2", "h3", "h4", "tr"):
                self.parts.append("\n")

        def handle_endtag(self, tag):
            if tag in ("script", "style"):
                self.skip = max(0, self.skip - 1)

        def handle_data(self, data):
            if not self.skip:
                self.parts.append(data)

    parser = _TextExtractor()
    parser.feed(html)
    text = "".join(parser.parts)
    return html_lib.unescape(re.sub(r"\n{2,}", "\n", text)).strip()


def _image_width_for(name: str) -> Inches:
    return Inches(DEFAULT_IMAGE_WIDTHS.get(name, 5.5))


def replace_image_markers(doc: Document, provider) -> int:
    """Replace [img:NAME] markers with inline pictures.

    `provider(name)` returns an ObjectData (image bytes + content type) or None.
    When a marker is alone in its paragraph the paragraph is cleared and the
    image becomes the paragraph content (centered alignment is preserved).
    Markers embedded in a run of text keep their surrounding text.
    """
    replaced = 0

    def _process(paragraph) -> None:
        nonlocal replaced
        text = paragraph.text
        if not IMAGE_MARKER_RE.search(text):
            return
        parts = IMAGE_MARKER_RE.split(text)  # [text, name, text, name, ...]
        paragraph.clear()  # preserves paragraph style/alignment
        for index, part in enumerate(parts):
            if index % 2 == 0:
                if part:
                    paragraph.add_run(part)
                continue
            image = provider(part)
            if image is None:
                continue  # marker without an image is dropped
            run = paragraph.add_run()
            run.add_picture(io.BytesIO(image.data), width=_image_width_for(part))
            replaced += 1

    for paragraph in doc.paragraphs:
        _process(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _process(paragraph)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                _process(paragraph)
    return replaced


def render_report(
    template_bytes: bytes, context: dict, image_provider
) -> bytes:
    """Render a template with the given context and image provider.

    Returns the editable DOCX bytes.
    """
    tpl = DocxTemplate(io.BytesIO(template_bytes))
    tpl.render(context)

    buf = io.BytesIO()
    tpl.save(buf)
    buf.seek(0)

    doc = Document(buf)
    replace_image_markers(doc, image_provider)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


def docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert DOCX bytes to PDF bytes using LibreOffice headless."""
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "report.docx")
        pdf_path = os.path.join(tmp, "report.pdf")
        Path(docx_path).write_bytes(docx_bytes)
        result = subprocess.run(
            [
                "soffice",
                "--headless",
                "--norestore",
                "--convert-to",
                "pdf",
                "--outdir",
                tmp,
                docx_path,
            ],
            capture_output=True,
            timeout=120,
            check=True,
        )
        if not os.path.exists(pdf_path):
            raise RuntimeError(
                "LibreOffice did not produce a PDF; stdout="
                + (result.stdout.decode(errors="ignore") or "")
            )
        return Path(pdf_path).read_bytes()


def docx_to_html(docx_bytes: bytes) -> str:
    """Convert DOCX bytes to HTML for browser preview (mammoth)."""
    import mammoth

    result = mammoth.convert_to_html(io.BytesIO(docx_bytes))
    if result.messages:
        # Non-fatal warnings only; embedded images are handled separately.
        pass
    return result.value


def image_provider_from(assets: dict[str, ObjectData]):
    """Build a provider closure that returns stored asset images by name."""

    def _provider(name: str) -> ObjectData | None:
        return assets.get(name)

    return _provider