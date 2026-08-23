"""Programmatic construction of the bundled sample NGO report template.

The resulting .docx contains jinja2 placeholders ({{ var }}) filled by docxtpl
and [img:NAME] markers swapped for images by the generation pipeline. It is
reproducible from code so the scaffold works out-of-the-box without Word.
"""

import io

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

TEAL = RGBColor(0x0B, 0x6E, 0x6B)
DARK = RGBColor(0x22, 0x33, 0x33)
GREY = RGBColor(0x6B, 0x72, 0x80)


def _style_base(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15


def _heading(doc: Document, text: str, size: int = 16) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = TEAL
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p


def _cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()

    logo = doc.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo.add_run("[img:logo]")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("{{ org_name }}")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = TEAL

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Annual Report {{ report_year }}")
    run.font.size = Pt(18)
    run.font.color.rgb = GREY

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tag.add_run("{{ tagline }}")
    run.italic = True
    run.font.size = Pt(12)

    doc.add_page_break()


def _mission(doc: Document) -> None:
    _heading(doc, "Our Mission")
    doc.add_paragraph("{{ mission.statement }}")


def _impact(doc: Document) -> None:
    _heading(doc, "Impact in Numbers")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = table.rows[0].cells
    for index, text in enumerate(("Beneficiaries served", "Communities reached", "Volunteers engaged")):
        cell = headers[index]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    row = table.add_row().cells
    values = ("{{ impact.beneficiaries }}", "{{ impact.communities }}", "{{ impact.volunteers }}")
    for index, value in enumerate(values):
        p = row[index].paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(14)
        run.font.color.rgb = TEAL
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _financials(doc: Document) -> None:
    _heading(doc, "Financial Highlights")
    chart = doc.add_paragraph()
    chart.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chart.add_run("[img:chart_funding]")
    doc.add_paragraph("Total funding: {{ financial.total }}")


def _donors(doc: Document) -> None:
    _heading(doc, "Donor Acknowledgment")
    doc.add_paragraph("{{ donors.acknowledgment }}")


def _goals(doc: Document) -> None:
    _heading(doc, "Looking Ahead")
    doc.add_paragraph("{{ future_goals }}")


def build_sample_template() -> bytes:
    doc = Document()
    _style_base(doc)
    _cover(doc)
    _mission(doc)
    _impact(doc)
    _financials(doc)
    _donors(doc)
    _goals(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


SAMPLE_SCHEMA = {
    "title": "NGO Annual Report",
    "description": "Standard annual report template with cover, mission, impact, financials, donor acknowledgment and future goals.",
    "sections": [
        {"key": "cover", "label": "Cover", "sort": 0},
        {"key": "mission", "label": "Mission", "sort": 1},
        {"key": "impact", "label": "Impact", "sort": 2},
        {"key": "financials", "label": "Financials", "sort": 3},
        {"key": "donors", "label": "Donor Acknowledgment", "sort": 4},
        {"key": "goals", "label": "Future Goals", "sort": 5},
    ],
    "section_map": {
        "mission": "mission.statement",
        "donors": "donors.acknowledgment",
        "goals": "future_goals",
    },
    "fields": [
        {
            "group": "Cover",
            "fields": [
                {"name": "org_name", "label": "Organization name", "type": "text", "path": "org_name", "required": True},
                {"name": "report_year", "label": "Report year", "type": "text", "path": "report_year", "required": True},
                {"name": "tagline", "label": "Tagline", "type": "textarea", "path": "tagline", "required": False},
                {"name": "logo", "label": "Logo image", "type": "image", "path": "logo", "placeholder": "logo", "required": False},
            ],
        },
        {
            "group": "Mission",
            "fields": [
                {"name": "mission_statement", "label": "Mission statement", "type": "textarea", "path": "mission.statement", "required": True},
            ],
        },
        {
            "group": "Impact",
            "fields": [
                {"name": "impact_beneficiaries", "label": "Beneficiaries served", "type": "number", "path": "impact.beneficiaries", "required": True},
                {"name": "impact_communities", "label": "Communities reached", "type": "number", "path": "impact.communities", "required": True},
                {"name": "impact_volunteers", "label": "Volunteers engaged", "type": "number", "path": "impact.volunteers", "required": True},
            ],
        },
        {
            "group": "Financials",
            "fields": [
                {"name": "financial_total", "label": "Total funding", "type": "number", "path": "financial.total", "required": True},
                {"name": "chart_funding", "label": "Funding chart image", "type": "image", "path": "chart_funding", "placeholder": "chart_funding", "required": False},
            ],
        },
        {
            "group": "Donor Acknowledgment",
            "fields": [
                {"name": "donors_ack", "label": "Donor acknowledgment text", "type": "textarea", "path": "donors.acknowledgment", "required": False},
            ],
        },
        {
            "group": "Future Goals",
            "fields": [
                {"name": "future_goals", "label": "Future goals", "type": "textarea", "path": "future_goals", "required": False},
            ],
        },
    ],
}


# ---------------------------------------------------------------------------
# Publication-style template — mirrors the layout of the five UN/NGO report
# samples (WHO Bhutan, UNICEF CPD, UNICEF donor brochure, UNDP biodiversity
# booklet, UN Women WEE brochure):
#   cover -> about/copyright -> contents -> foreword -> quote page ->
#   executive overview (stat callouts) -> impact (chart + pull quote) ->
#   programmes (photo blocks) -> milestones timeline -> financials (budget
#   table + chart) -> donor acknowledgment (pull quote) -> looking ahead ->
#   closing statement -> back cover.
# ---------------------------------------------------------------------------

A4_W, A4_H = Inches(8.27), Inches(11.69)


def _page_break(doc: Document) -> None:
    doc.add_page_break()


def _cover_page(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    logo = doc.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo.add_run("[img:logo]")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("{{ org_name }}")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = TEAL

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Annual Report {{ report_year }}")
    run.font.size = Pt(18)
    run.font.color.rgb = GREY

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tag.add_run("{{ tagline }}")
    run.italic = True
    run.font.size = Pt(13)

    doc.add_paragraph()
    cover_line = doc.add_paragraph()
    cover_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_line.add_run("{{ report_type }}")
    run.font.size = Pt(11)
    run.font.color.rgb = GREY


def _about_page(doc: Document) -> None:
    _heading(doc, "About this report", 20)
    doc.add_paragraph("{{ about_report }}")
    doc.add_paragraph()
    doc.add_paragraph("© {{ report_year }} {{ org_name }}. All rights reserved.")
    doc.add_paragraph("Published {{ report_year }}. Cover image © {{ org_name }}.")


def _contents_page(doc: Document) -> None:
    _heading(doc, "Contents", 20)
    items = [
        ("Foreword", "1"),
        ("Executive Overview", "3"),
        ("Impact & Results", "5"),
        ("Our Programmes", "7"),
        ("Milestones", "11"),
        ("Financial Highlights", "13"),
        ("Donor Acknowledgment", "15"),
        ("Looking Ahead", "16"),
        ("Closing Statement", "17"),
    ]
    for label, page in items:
        p = doc.add_paragraph()
        run = p.add_run(f"{label:<40}{page}")
        run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(10)


def _foreword_page(doc: Document) -> None:
    _heading(doc, "Foreword", 20)
    doc.add_paragraph("{{ foreword }}")
    doc.add_paragraph()
    sig = doc.add_paragraph()
    run = sig.add_run("{{ leader_name }}")
    run.bold = True
    run.font.size = Pt(12)
    title = doc.add_paragraph()
    run = title.add_run("{{ leader_title }}")
    run.font.color.rgb = GREY
    run.font.size = Pt(11)


def _quote_page(doc: Document) -> None:
    for _ in range(6):
        doc.add_paragraph()
    q = doc.add_paragraph()
    q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = q.add_run("“{{ opening_quote }}”")
    run.italic = True
    run.font.size = Pt(22)
    run.font.color.rgb = TEAL
    attr = doc.add_paragraph()
    attr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = attr.add_run("— {{ opening_quote_author }}")
    run.font.size = Pt(12)
    run.font.color.rgb = GREY


def _stat_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(rows))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    for index, (label, _value) in enumerate(rows):
        p = header[index].paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    row = table.add_row().cells
    for index, (_label, value) in enumerate(rows):
        p = row[index].paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(20)
        run.font.color.rgb = TEAL
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _pull_quote(doc: Document, quote_var: str, author_var: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"“{{{quote_var}}}”")
    run.italic = True
    run.font.size = Pt(15)
    run.font.color.rgb = GREY
    a = doc.add_paragraph()
    run = a.add_run(f"— {{{author_var}}}")
    run.font.size = Pt(11)
    run.font.color.rgb = GREY


def _overview_page(doc: Document) -> None:
    _heading(doc, "Executive Overview", 20)
    doc.add_paragraph("{{ about_intro }}")
    doc.add_paragraph()
    _stat_table(
        doc,
        [
            ("Beneficiaries served", "{{ impact.beneficiaries }}"),
            ("Communities reached", "{{ impact.communities }}"),
            ("Volunteers engaged", "{{ impact.volunteers }}"),
            ("Districts served", "{{ impact.districts }}"),
        ],
    )
    doc.add_paragraph()
    doc.add_paragraph("{{ overview_narrative }}")


def _impact_page(doc: Document) -> None:
    _heading(doc, "Impact & Results", 20)
    doc.add_paragraph("{{ impact_summary }}")
    chart = doc.add_paragraph()
    chart.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chart.add_run("[img:chart_impact]")
    doc.add_paragraph()
    _pull_quote(doc, "impact_quote", "impact_quote_author")


def _programmes_page(doc: Document) -> None:
    _heading(doc, "Our Programmes", 20)
    doc.add_paragraph("{{ programs_intro }}")
    for index in range(1, 5):
        photo = doc.add_paragraph()
        photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        photo.add_run(f"[img:program_{index}]")
        p = doc.add_paragraph()
        run = p.add_run(f"{{{{ program_{index}_name }}}}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = TEAL
        doc.add_paragraph(f"{{{{ program_{index}_desc }}}}")


def _milestones_page(doc: Document) -> None:
    _heading(doc, "Milestones", 20)
    doc.add_paragraph("{{ milestones_intro }}")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].paragraphs[0].add_run("Year").bold = True
    header[1].paragraphs[0].add_run("Milestone").bold = True
    for index in range(1, 6):
        row = table.add_row().cells
        row[0].paragraphs[0].add_run(f"{{{{ milestone_{index}_year }}}}").bold = True
        row[1].paragraphs[0].add_run(f"{{{{ milestone_{index}_text }}}}")


def _financials_page(doc: Document) -> None:
    _heading(doc, "Financial Highlights", 20)
    doc.add_paragraph("{{ financial_summary }}")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header = table.rows[0].cells
    for index, label in enumerate(("Category", "Amount", "Share")):
        header[index].paragraphs[0].add_run(label).bold = True
    rows = [
        ("Programmes & operations", "{{ financial.programmes }}", "{{ financial.programmes_share }}"),
        ("Fundraising & administration", "{{ financial.admin }}", "{{ financial.admin_share }}"),
        ("Total", "{{ financial.total }}", "100%"),
    ]
    for label, amount, share in rows:
        row = table.add_row().cells
        row[0].paragraphs[0].add_run(label)
        row[1].paragraphs[0].add_run(amount)
        row[2].paragraphs[0].add_run(share)
    doc.add_paragraph()
    chart = doc.add_paragraph()
    chart.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chart.add_run("[img:chart_funding]")


def _donors_page(doc: Document) -> None:
    _heading(doc, "Donor Acknowledgment", 20)
    doc.add_paragraph("{{ donors_ack }}")
    doc.add_paragraph()
    _pull_quote(doc, "donor_quote", "donor_quote_author")


def _goals_page(doc: Document) -> None:
    _heading(doc, "Looking Ahead", 20)
    doc.add_paragraph("{{ future_goals }}")


def _closing_page(doc: Document) -> None:
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("{{ closing_statement }}")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = TEAL
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("— {{ org_name }}")
    run.font.size = Pt(13)
    run.font.color.rgb = GREY


def _back_cover(doc: Document) -> None:
    _heading(doc, "Contact", 20)
    lines = [
        "{{ contact_address }}",
        "{{ contact_phone }}",
        "{{ contact_email }}",
        "{{ contact_website }}",
        "{{ contact_social }}",
    ]
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.color.rgb = GREY


def build_publication_template() -> bytes:
    doc = Document()
    _style_base(doc)

    section = doc.sections[0]
    section.page_width = A4_W
    section.page_height = A4_H
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    _cover_page(doc)
    _page_break(doc)
    _about_page(doc)
    _page_break(doc)
    _contents_page(doc)
    _page_break(doc)
    _foreword_page(doc)
    _page_break(doc)
    _quote_page(doc)
    _page_break(doc)
    _overview_page(doc)
    _page_break(doc)
    _impact_page(doc)
    _page_break(doc)
    _programmes_page(doc)
    _page_break(doc)
    _milestones_page(doc)
    _page_break(doc)
    _financials_page(doc)
    _page_break(doc)
    _donors_page(doc)
    _page_break(doc)
    _goals_page(doc)
    _page_break(doc)
    _closing_page(doc)
    _page_break(doc)
    _back_cover(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


PUBLICATION_SCHEMA = {
    "title": "Annual Report",
    "description": "Brochure/book-style annual report mirroring UN/NGO publication layouts: cover, foreword, quote page, executive overview with stat callouts, impact & results, programmes, milestones timeline, financials, donor acknowledgment, looking ahead, closing statement and back cover.",
    "sections": [
        {"key": "foreword", "label": "Foreword", "sort": 1},
        {"key": "overview", "label": "Executive Overview", "sort": 2},
        {"key": "impact", "label": "Impact & Results", "sort": 3},
        {"key": "programs", "label": "Our Programmes", "sort": 4},
        {"key": "milestones", "label": "Milestones", "sort": 5},
        {"key": "financials", "label": "Financial Highlights", "sort": 6},
        {"key": "donors", "label": "Donor Acknowledgment", "sort": 7},
        {"key": "goals", "label": "Looking Ahead", "sort": 8},
        {"key": "closing", "label": "Closing Statement", "sort": 9},
    ],
    "section_map": {
        "foreword": "foreword",
        "overview": "about_intro",
        "impact": "impact_summary",
        "programs": "programs_intro",
        "donors": "donors_ack",
        "goals": "future_goals",
        "closing": "closing_statement",
    },
    "fields": [
        {
            "group": "Cover",
            "fields": [
                {"name": "org_name", "label": "Organization name", "type": "text", "path": "org_name", "required": True, "placeholder": "e.g. BrightPath Foundation"},
                {"name": "report_year", "label": "Report year", "type": "text", "path": "report_year", "required": True, "placeholder": "e.g. 2025"},
                {"name": "tagline", "label": "Tagline", "type": "textarea", "path": "tagline", "required": False, "placeholder": "A short tagline — e.g. \u201cBuilding healthier communities together\u201d"},
                {"name": "report_type", "label": "Report subtitle", "type": "text", "path": "report_type", "required": False, "placeholder": "e.g. Annual Report of Activities and Impact"},
                {"name": "logo", "label": "Logo image", "type": "image", "path": "logo", "placeholder": "logo", "required": False},
            ],
        },
        {
            "group": "About & Foreword",
            "fields": [
                {"name": "about_report", "label": "About this report", "type": "textarea", "path": "about_report", "required": False, "placeholder": "A short description of this report and its purpose…"},
                {"name": "foreword", "label": "Foreword text", "type": "textarea", "path": "foreword", "required": True, "placeholder": "A letter from leadership introducing the year…"},
                {"name": "leader_name", "label": "Signed by (name)", "type": "text", "path": "leader_name", "required": False, "placeholder": "e.g. Dr. A. Sharma"},
                {"name": "leader_title", "label": "Signed by (title)", "type": "text", "path": "leader_title", "required": False, "placeholder": "e.g. Executive Director"},
            ],
        },
        {
            "group": "Opening Quote",
            "fields": [
                {"name": "opening_quote", "label": "Opening quote", "type": "textarea", "path": "opening_quote", "required": False, "placeholder": "An inspiring quote that frames the year…"},
                {"name": "opening_quote_author", "label": "Quote attribution", "type": "text", "path": "opening_quote_author", "required": False, "placeholder": "e.g. Founder, {org}"},
            ],
        },
        {
            "group": "Executive Overview",
            "fields": [
                {"name": "about_intro", "label": "Overview introduction", "type": "textarea", "path": "about_intro", "required": True, "placeholder": "Who you are and what the report covers…"},
                {"name": "overview_narrative", "label": "Overview narrative", "type": "textarea", "path": "overview_narrative", "required": False, "placeholder": "The story of the year in 2–3 paragraphs…"},
                {"name": "impact_beneficiaries", "label": "Beneficiaries served", "type": "number", "path": "impact.beneficiaries", "required": True, "placeholder": "e.g. 12,800"},
                {"name": "impact_communities", "label": "Communities reached", "type": "number", "path": "impact.communities", "required": True, "placeholder": "e.g. 42"},
                {"name": "impact_volunteers", "label": "Volunteers engaged", "type": "number", "path": "impact.volunteers", "required": True, "placeholder": "e.g. 350"},
                {"name": "impact_districts", "label": "Districts served", "type": "number", "path": "impact.districts", "required": False, "placeholder": "e.g. 15"},
            ],
        },
        {
            "group": "Impact & Results",
            "fields": [
                {"name": "impact_summary", "label": "Impact summary", "type": "textarea", "path": "impact_summary", "required": True, "placeholder": "Key results achieved this year…"},
                {"name": "impact_quote", "label": "Impact pull quote", "type": "textarea", "path": "impact_quote", "required": False, "placeholder": "A short quote capturing impact…"},
                {"name": "impact_quote_author", "label": "Impact quote attribution", "type": "text", "path": "impact_quote_author", "required": False, "placeholder": "e.g. Programme Manager"},
                {"name": "chart_impact", "label": "Impact chart image", "type": "image", "path": "chart_impact", "placeholder": "chart_impact", "required": False},
            ],
        },
        {
            "group": "Our Programmes",
            "fields": [
                {"name": "programs_intro", "label": "Programmes introduction", "type": "textarea", "path": "programs_intro", "required": False, "placeholder": "An overview of your programmes…"},
                {"name": "program_1_name", "label": "Programme 1 name", "type": "text", "path": "program_1_name", "required": False, "placeholder": "e.g. Clean Water Initiative"},
                {"name": "program_1_desc", "label": "Programme 1 description", "type": "textarea", "path": "program_1_desc", "required": False, "placeholder": "What the programme does and its outcomes…"},
                {"name": "program_2_name", "label": "Programme 2 name", "type": "text", "path": "program_2_name", "required": False, "placeholder": "e.g. Education for All"},
                {"name": "program_2_desc", "label": "Programme 2 description", "type": "textarea", "path": "program_2_desc", "required": False, "placeholder": "What the programme does and its outcomes…"},
                {"name": "program_3_name", "label": "Programme 3 name", "type": "text", "path": "program_3_name", "required": False, "placeholder": "e.g. Nutrition Support"},
                {"name": "program_3_desc", "label": "Programme 3 description", "type": "textarea", "path": "program_3_desc", "required": False, "placeholder": "What the programme does and its outcomes…"},
                {"name": "program_4_name", "label": "Programme 4 name", "type": "text", "path": "program_4_name", "required": False, "placeholder": "e.g. Youth Leadership"},
                {"name": "program_4_desc", "label": "Programme 4 description", "type": "textarea", "path": "program_4_desc", "required": False, "placeholder": "What the programme does and its outcomes…"},
                {"name": "program_1_img", "label": "Programme 1 photo", "type": "image", "path": "program_1", "placeholder": "program_1", "required": False},
                {"name": "program_2_img", "label": "Programme 2 photo", "type": "image", "path": "program_2", "placeholder": "program_2", "required": False},
                {"name": "program_3_img", "label": "Programme 3 photo", "type": "image", "path": "program_3", "placeholder": "program_3", "required": False},
                {"name": "program_4_img", "label": "Programme 4 photo", "type": "image", "path": "program_4", "placeholder": "program_4", "required": False},
            ],
        },
        {
            "group": "Milestones",
            "fields": [
                {"name": "milestones_intro", "label": "Milestones introduction", "type": "textarea", "path": "milestones_intro", "required": False, "placeholder": "How the journey unfolded…"},
                {"name": "milestone_1_year", "label": "Milestone 1 year", "type": "text", "path": "milestone_1_year", "required": False, "placeholder": "e.g. 2019"},
                {"name": "milestone_1_text", "label": "Milestone 1", "type": "text", "path": "milestone_1_text", "required": False, "placeholder": "e.g. Founded and first outreach"},
                {"name": "milestone_2_year", "label": "Milestone 2 year", "type": "text", "path": "milestone_2_year", "required": False, "placeholder": "e.g. 2021"},
                {"name": "milestone_2_text", "label": "Milestone 2", "type": "text", "path": "milestone_2_text", "required": False, "placeholder": "e.g. Reached 5,000 beneficiaries"},
                {"name": "milestone_3_year", "label": "Milestone 3 year", "type": "text", "path": "milestone_3_year", "required": False, "placeholder": "e.g. 2023"},
                {"name": "milestone_3_text", "label": "Milestone 3", "type": "text", "path": "milestone_3_text", "required": False, "placeholder": "e.g. Expanded to 10 districts"},
                {"name": "milestone_4_year", "label": "Milestone 4 year", "type": "text", "path": "milestone_4_year", "required": False, "placeholder": "e.g. 2024"},
                {"name": "milestone_4_text", "label": "Milestone 4", "type": "text", "path": "milestone_4_text", "required": False, "placeholder": "e.g. First 100% digital program"},
                {"name": "milestone_5_year", "label": "Milestone 5 year", "type": "text", "path": "milestone_5_year", "required": False, "placeholder": "e.g. 2025"},
                {"name": "milestone_5_text", "label": "Milestone 5", "type": "text", "path": "milestone_5_text", "required": False, "placeholder": "e.g. 1 million people reached"},
            ],
        },
        {
            "group": "Financial Highlights",
            "fields": [
                {"name": "financial_summary", "label": "Financial summary", "type": "textarea", "path": "financial_summary", "required": False, "placeholder": "A short narrative of the year's finances…"},
                {"name": "financial_programmes", "label": "Programmes & operations amount", "type": "text", "path": "financial.programmes", "required": False, "placeholder": "e.g. $1,100,000"},
                {"name": "financial_programmes_share", "label": "Programmes share", "type": "text", "path": "financial.programmes_share", "required": False, "placeholder": "e.g. 78%"},
                {"name": "financial_admin", "label": "Fundraising & admin amount", "type": "text", "path": "financial.admin", "required": False, "placeholder": "e.g. $300,000"},
                {"name": "financial_admin_share", "label": "Fundraising & admin share", "type": "text", "path": "financial.admin_share", "required": False, "placeholder": "e.g. 22%"},
                {"name": "financial_total", "label": "Total funding", "type": "text", "path": "financial.total", "required": True, "placeholder": "e.g. $1,400,000"},
                {"name": "chart_funding", "label": "Funding chart image", "type": "image", "path": "chart_funding", "placeholder": "chart_funding", "required": False},
            ],
        },
        {
            "group": "Donor Acknowledgment",
            "fields": [
                {"name": "donors_ack", "label": "Donor acknowledgment text", "type": "textarea", "path": "donors_ack", "required": False, "placeholder": "Thank your donors and partners here…"},
                {"name": "donor_quote", "label": "Donor pull quote", "type": "textarea", "path": "donor_quote", "required": False, "placeholder": "A quote from a donor or partner…"},
                {"name": "donor_quote_author", "label": "Donor quote attribution", "type": "text", "path": "donor_quote_author", "required": False, "placeholder": "e.g. Major Donor"},
            ],
        },
        {
            "group": "Looking Ahead & Closing",
            "fields": [
                {"name": "future_goals", "label": "Future goals", "type": "textarea", "path": "future_goals", "required": False, "placeholder": "Goals for the coming year…"},
                {"name": "closing_statement", "label": "Closing statement", "type": "textarea", "path": "closing_statement", "required": False, "placeholder": "A bold closing line, e.g. \u201cTogether, we can\u2026\u201d"},
            ],
        },
        {
            "group": "Contact (back cover)",
            "fields": [
                {"name": "contact_address", "label": "Address", "type": "text", "path": "contact_address", "required": False, "placeholder": "e.g. 12 Community Lane, New Delhi 110 001"},
                {"name": "contact_phone", "label": "Phone", "type": "text", "path": "contact_phone", "required": False, "placeholder": "e.g. +91 11 1234 5678"},
                {"name": "contact_email", "label": "Email", "type": "text", "path": "contact_email", "required": False, "placeholder": "e.g. hello@yourorg.org"},
                {"name": "contact_website", "label": "Website", "type": "text", "path": "contact_website", "required": False, "placeholder": "e.g. www.yourorg.org"},
                {"name": "contact_social", "label": "Social links", "type": "text", "path": "contact_social", "required": False, "placeholder": "e.g. @yourorg on X, LinkedIn, Instagram"},
            ],
        },
    ],
}


# ---------------------------------------------------------------------------
# Magazine-style template (matches the described 5-page preview):

from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

BLUE = RGBColor(0x44, 0x84, 0xCE)
BLUE_HEX = "4484CE"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _shade_cell(cell, hex_fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _fixed_table(doc, widths: list[float]):
    table = doc.add_table(rows=1, cols=len(widths))
    table.autofit = False
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)
    for cell, width in zip(table.rows[0].cells, widths):
        cell.width = Inches(width)
    return table


def _white_run(paragraph, text: str, size: int = 11, bold: bool = False):
    run = paragraph.add_run(text)
    run.font.color.rgb = WHITE
    run.font.size = Pt(size)
    run.bold = bold
    return run


def _blue_para(cell, text: str, size: int = 11, bold: bool = False, space: int = 5):
    p = cell.add_paragraph()
    _white_run(p, text, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space)
    return p


def _photo_cell(cell, marker: str) -> None:
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(marker)


def _section_header(cell, title: str, authors: str | None = None) -> None:
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = BLUE
    if authors:
        a = cell.add_paragraph()
        run = a.add_run(authors)
        run.font.color.rgb = GREY
        run.font.size = Pt(9)
        a.paragraph_format.space_after = Pt(8)


def _circle_number(cell, text: str) -> None:
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = BLUE
    p.paragraph_format.space_after = Pt(6)


def _toc_item(cell, label: str, authors: str) -> None:
    p = cell.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    a = cell.add_paragraph()
    run = a.add_run(authors)
    run.font.color.rgb = GREY
    run.font.size = Pt(9)
    a.paragraph_format.space_after = Pt(8)


def build_magazine_template() -> bytes:
    doc = Document()
    _style_base(doc)
    section = doc.sections[0]
    section.page_width = A4_W
    section.page_height = A4_H
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)

    # Footer with a page number (right aligned, small grey).
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    f_run = footer_p.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    f_run._r.append(fld1)
    f_run._r.append(instr)
    f_run._r.append(fld2)
    f_run.font.size = Pt(9)
    f_run.font.color.rgb = GREY

    # ---- Page 1: Cover ----
    cover = _fixed_table(doc, [2.7, 6.4])
    left, right = cover.rows[0].cells
    _shade_cell(right, BLUE_HEX)
    _photo_cell(left, "[img:cover_photo]")
    logo_p = right.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_p.add_run("[img:cover_logo]")
    for _ in range(3):
        right.add_paragraph()
    year_p = right.add_paragraph()
    year_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _white_run(year_p, "{{ report_year }}", size=40, bold=True)
    title_p = right.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _white_run(title_p, "ANNUAL REPORT", size=30, bold=True)
    tag = right.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _white_run(tag, "{{ tagline }}", size=13)
    for _ in range(6):
        right.add_paragraph()
    authors = right.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _white_run(authors, "{{ cover_authors }}", size=11)
    doc.add_page_break()

    # ---- Page 2: Contents & Introduction ----
    contents = _fixed_table(doc, [2.7, 6.4])
    left, right = contents.rows[0].cells
    _photo_cell(left, "[img:contents_photo]")
    for num in ("1", "2", "3"):
        _circle_number(left, num)
    h = right.paragraphs[0]
    run = h.add_run("CONTENTS")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = BLUE
    pPr = h._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), BLUE_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)
    _toc_item(right, "1 STRATEGY", "KALAN ABHEEDI, ROYCE CHALMER")
    _toc_item(right, "2 FINANCE", "AMANDA SULLY, JANN VAN HUYSEN")
    _toc_item(right, "3 PROJECTS", "JEANNETTE MOSS, ELENA SONG")
    band = _fixed_table(doc, [9.1])
    cell = band.rows[0].cells[0]
    _shade_cell(cell, BLUE_HEX)
    _blue_para(cell, "{{ intro_title }}", size=18, bold=True)
    _blue_para(cell, "{{ intro_text }}", size=11)
    _blue_para(cell, "{{ intro_extra }}", size=10)
    doc.add_page_break()

    # ---- Page 3: Strategy ----
    strategy = _fixed_table(doc, [2.7, 6.4])
    left, right = strategy.rows[0].cells
    _photo_cell(left, "[img:strategy_photo]")
    _section_header(right, "STRATEGY", "KALAN ABHEEDI, ROYCE CHALMER")
    stat = right.add_paragraph()
    run = stat.add_run("{{ strategy_funds_2029 }}")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = BLUE
    src = right.add_paragraph()
    run = src.add_run("from {{ strategy_funds_2028 }}")
    run.font.color.rgb = GREY
    note = right.add_paragraph()
    run = note.add_run("(funds raised annually in {{ report_year }} vs. last year)")
    run.font.color.rgb = GREY
    run.font.size = Pt(9)
    stat = right.add_paragraph()
    run = stat.add_run("{{ strategy_people_2029 }}")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = BLUE
    src = right.add_paragraph()
    run = src.add_run("from {{ strategy_people_2028 }}")
    run.font.color.rgb = GREY
    note = right.add_paragraph()
    run = note.add_run("(people served annually {{ report_year }} vs. last year)")
    run.font.color.rgb = GREY
    run.font.size = Pt(9)
    band = _fixed_table(doc, [9.1])
    cell = band.rows[0].cells[0]
    _shade_cell(cell, BLUE_HEX)
    _blue_para(cell, "Looking Back at {{ report_year }}...", size=16, bold=True)
    _blue_para(cell, "OUTREACH", size=12, bold=True)
    _blue_para(cell, "{{ outreach_emails }} EMAILS", size=11)
    _blue_para(cell, "{{ outreach_conversations }} CONVERSATIONS", size=11)
    _blue_para(cell, "{{ outreach_speeches }} SPEECHES", size=11)
    _blue_para(cell, "VOLUNTEER RETENTION", size=12, bold=True)
    _blue_para(cell, "{{ volunteer_growth }} GROWTH", size=11)
    _blue_para(cell, "{{ volunteer_hours }} HOURS", size=11)
    _blue_para(cell, "{{ volunteer_party }}", size=11)
    _blue_para(cell, "AWARENESS", size=12, bold=True)
    _blue_para(cell, "DIGITAL MARKETING", size=11)
    _blue_para(cell, "GUERILLA MARKETING", size=11)
    _blue_para(cell, "CAMPAIGN ADVERTISING", size=11)
    doc.add_page_break()

    # ---- Page 4: Finance ----
    finance = _fixed_table(doc, [2.7, 6.4])
    left, right = finance.rows[0].cells
    _photo_cell(left, "[img:finance_photo]")
    _section_header(right, "FINANCE", "AMANDA SULLY, JANN VAN HUYSEN")
    src = right.add_paragraph()
    run = src.add_run("Sources of Funding")
    run.bold = True
    run.font.size = Pt(12)
    src = right.add_paragraph()
    src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src.add_run("[img:finance_sources]")
    src = right.add_paragraph()
    run = src.add_run("Annual Expenses")
    run.bold = True
    run.font.size = Pt(12)
    src = right.add_paragraph()
    src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    src.add_run("[img:finance_expenses]")
    band = _fixed_table(doc, [9.1])
    cell = band.rows[0].cells[0]
    _shade_cell(cell, BLUE_HEX)
    _blue_para(cell, "Additional Remarks", size=16, bold=True)
    _blue_para(cell, "EVENTS: {{ finance_events }}", size=11)
    _blue_para(cell, "CONFERENCES: {{ finance_conferences }}", size=11)
    _blue_para(cell, "DIGITAL MARKETING: {{ finance_digital }}", size=11)
    doc.add_page_break()

    # ---- Page 5: Projects ----
    projects = _fixed_table(doc, [2.7, 6.4])
    left, right = projects.rows[0].cells
    _photo_cell(left, "[img:projects_photo]")
    _section_header(right, "PROJECTS", "JEANNETTE MOSS, ELENA SONG")
    h = right.add_paragraph()
    run = h.add_run("Annual Fundraiser")
    run.bold = True
    run.font.size = Pt(12)
    right.add_paragraph("{{ projects_fundraiser }}")
    h = right.add_paragraph()
    run = h.add_run("#itsnottheflu")
    run.bold = True
    run.font.size = Pt(12)
    right.add_paragraph("{{ projects_campaign }}")
    band = _fixed_table(doc, [9.1])
    cell = band.rows[0].cells[0]
    _shade_cell(cell, BLUE_HEX)
    _blue_para(cell, "{{ projects_handwash_title }}", size=16, bold=True)
    _blue_para(cell, "{{ projects_handwash_text }}", size=11)
    _blue_para(cell, "{{ projects_checklist }}", size=11)
    photo2 = cell.add_paragraph()
    photo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    photo2.add_run("[img:projects_photo2]")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


MAGAZINE_SCHEMA = {
    "title": "Annual Report",
    "description": "Magazine-style annual report: split cover with photo and solid blue title block, contents & introduction, strategy stats page, finance page with charts, and projects page.",
    "sections": [
        {"key": "contents", "label": "Contents & Introduction", "sort": 1},
        {"key": "strategy", "label": "Strategy", "sort": 2},
        {"key": "finance", "label": "Finance", "sort": 3},
        {"key": "projects", "label": "Projects", "sort": 4},
    ],
    "section_map": {
        "contents": "intro_text",
        "strategy": "strategy_notes",
        "finance": "finance_events",
        "projects": "projects_fundraiser",
    },
    "fields": [
        {
            "group": "Cover",
            "fields": [
                {"name": "org_name", "label": "Organization name", "type": "text", "path": "org_name", "required": True, "placeholder": "e.g. BrightPath Foundation"},
                {"name": "report_year", "label": "Report year", "type": "text", "path": "report_year", "required": True, "placeholder": "e.g. 2029"},
                {"name": "tagline", "label": "Tagline", "type": "textarea", "path": "tagline", "required": False, "placeholder": "A short tagline shown under the title"},
                {"name": "cover_authors", "label": "Contributors (cover)", "type": "textarea", "path": "cover_authors", "required": False, "placeholder": "e.g. KALAN ABHEEDI, ROYCE CHALMER"},
                {"name": "cover_logo", "label": "Logo image", "type": "image", "path": "cover_logo", "placeholder": "cover_logo", "required": False},
                {"name": "cover_photo", "label": "Cover photo", "type": "image", "path": "cover_photo", "placeholder": "cover_photo", "required": False},
            ],
        },
        {
            "group": "Contents & Introduction",
            "fields": [
                {"name": "contents_photo", "label": "Contents photo", "type": "image", "path": "contents_photo", "placeholder": "contents_photo", "required": False},
                {"name": "intro_title", "label": "Introduction title", "type": "text", "path": "intro_title", "required": False, "placeholder": "e.g. The Road We Walked Together"},
                {"name": "intro_text", "label": "Introduction text", "type": "textarea", "path": "intro_text", "required": False, "placeholder": "A short welcome paragraph…"},
                {"name": "intro_extra", "label": "Introduction extra", "type": "textarea", "path": "intro_extra", "required": False, "placeholder": "Optional second paragraph…"},
            ],
        },
        {
            "group": "Strategy",
            "fields": [
                {"name": "strategy_photo", "label": "Strategy photo", "type": "image", "path": "strategy_photo", "placeholder": "strategy_photo", "required": False},
                {"name": "strategy_funds_2029", "label": "Funds raised (this year)", "type": "text", "path": "strategy_funds_2029", "required": False, "placeholder": "e.g. $14,500,200"},
                {"name": "strategy_funds_2028", "label": "Funds raised (last year)", "type": "text", "path": "strategy_funds_2028", "required": False, "placeholder": "e.g. $13,400,700"},
                {"name": "strategy_people_2029", "label": "People served (this year)", "type": "text", "path": "strategy_people_2029", "required": False, "placeholder": "e.g. 15,200"},
                {"name": "strategy_people_2028", "label": "People served (last year)", "type": "text", "path": "strategy_people_2028", "required": False, "placeholder": "e.g. 13,700"},
                {"name": "strategy_notes", "label": "Strategy narrative", "type": "textarea", "path": "strategy_notes", "required": False, "placeholder": "A short narrative for the strategy section…"},
                {"name": "outreach_emails", "label": "Emails sent", "type": "text", "path": "outreach_emails", "required": False, "placeholder": "e.g. 23,000"},
                {"name": "outreach_conversations", "label": "Conversations", "type": "text", "path": "outreach_conversations", "required": False, "placeholder": "e.g. 12,000"},
                {"name": "outreach_speeches", "label": "Speeches", "type": "text", "path": "outreach_speeches", "required": False, "placeholder": "e.g. 23"},
                {"name": "volunteer_growth", "label": "Volunteer growth", "type": "text", "path": "volunteer_growth", "required": False, "placeholder": "e.g. 13%"},
                {"name": "volunteer_hours", "label": "Volunteer hours", "type": "text", "path": "volunteer_hours", "required": False, "placeholder": "e.g. 460,000"},
                {"name": "volunteer_party", "label": "Year-end note", "type": "text", "path": "volunteer_party", "required": False, "placeholder": "e.g. 1 great year-end party"},
            ],
        },
        {
            "group": "Finance",
            "fields": [
                {"name": "finance_photo", "label": "Finance photo", "type": "image", "path": "finance_photo", "placeholder": "finance_photo", "required": False},
                {"name": "finance_sources", "label": "Sources-of-funding chart", "type": "image", "path": "finance_sources", "placeholder": "finance_sources", "required": False},
                {"name": "finance_expenses", "label": "Expenses chart", "type": "image", "path": "finance_expenses", "placeholder": "finance_expenses", "required": False},
                {"name": "finance_events", "label": "Events remark", "type": "textarea", "path": "finance_events", "required": False, "placeholder": "Planning fundraisers, community gatherings, sporting events…"},
                {"name": "finance_conferences", "label": "Conferences remark", "type": "textarea", "path": "finance_conferences", "required": False, "placeholder": "Attending conferences as guest or speaker…"},
                {"name": "finance_digital", "label": "Digital marketing remark", "type": "textarea", "path": "finance_digital", "required": False, "placeholder": "A digital roadmap to build our following…"},
            ],
        },
        {
            "group": "Projects",
            "fields": [
                {"name": "projects_photo", "label": "Projects photo", "type": "image", "path": "projects_photo", "placeholder": "projects_photo", "required": False},
                {"name": "projects_fundraiser", "label": "Annual fundraiser text", "type": "textarea", "path": "projects_fundraiser", "required": False, "placeholder": "The annual fundraiser is not the only event…"},
                {"name": "projects_campaign", "label": "#itsnottheflu text", "type": "textarea", "path": "projects_campaign", "required": False, "placeholder": "Inspired by our digital marketing team…"},
                {"name": "projects_handwash_title", "label": "Hand-wash campaign title", "type": "text", "path": "projects_handwash_title", "required": False, "placeholder": "If you're happy and you know it, wash your hands"},
                {"name": "projects_handwash_text", "label": "Hand-wash campaign text", "type": "textarea", "path": "projects_handwash_text", "required": False, "placeholder": "Meningitis can be transmitted through sharing germs…"},
                {"name": "projects_checklist", "label": "Campaign checklist", "type": "textarea", "path": "projects_checklist", "required": False, "placeholder": "• Awareness posters in urban areas.\n• Promotional video…\n• Public washing stations…"},
                {"name": "projects_photo2", "label": "Campaign photo", "type": "image", "path": "projects_photo2", "placeholder": "projects_photo2", "required": False},
            ],
        },
    ],
}


BUNDLED_TEMPLATES = [
    ("Annual Report", MAGAZINE_SCHEMA, build_magazine_template),
]